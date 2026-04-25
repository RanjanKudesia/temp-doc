"""DOCX extraction pipeline — chunking-optimised.

Changes vs. the full extraction pipeline:
- runs[] is NOT extracted (ChunkEngine only reads paragraph.text / style / list fields)
- cell.paragraphs is NOT built (ChunkEngine only reads cell.text)
- _extract_styles() removed — returns []
- _extract_document_defaults() removed — returns None
- json.dumps() removed from logging
- paragraph.alignment removed
- table.style removed
- len(inline_shapes) removed from structure log
- _resolve_list_formatting() cached by (num_id, ilvl)
- body_order built in a single iterchildren() pass alongside extraction
- Dead imports removed (base64, json, mimetypes, Run, RT, etree)
"""

import logging
import time
from io import BytesIO
from typing import Any
from zipfile import is_zipfile

from docx import Document
from docx.document import Document as DocumentObject
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

XML_NS = {
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
}

_W_VAL = "w:val"


class DocxExtractionPipeline:
    """Extract paragraphs and tables from DOCX files (chunking-optimised)."""

    def __init__(self) -> None:
        self.logger = logging.getLogger(__name__)

    def run(self, file_bytes: bytes, include_media: bool = True) -> dict[str, Any]:
        """Parse a DOCX byte stream and return extracted JSON payload."""
        t0 = time.perf_counter()
        self.logger.info(
            "DOCX extraction started",
            extra={"file_size_bytes": len(file_bytes)},
        )

        if not is_zipfile(BytesIO(file_bytes)):
            raise ValueError(
                "Invalid DOCX file: not a valid ZIP archive. File may be corrupted.")

        try:
            document = Document(BytesIO(file_bytes))
        except (ValueError, TypeError, OSError, KeyError) as e:
            raise ValueError(
                f"Failed to parse DOCX document: {str(e)}") from e

        extracted = self._extract_document(document)

        elapsed_ms = round((time.perf_counter() - t0) * 1000)
        self.logger.info(
            "DOCX extraction complete",
            extra={
                "elapsed_ms": elapsed_ms,
                "elapsed_s": round(elapsed_ms / 1000, 3),
                "paragraphs_extracted": len(extracted.get("paragraphs", [])),
                "tables_extracted": len(extracted.get("tables", [])),
            },
        )
        return extracted

    def _extract_document(self, document: DocumentObject) -> dict[str, Any]:
        """Extract all document content in a single body pass."""
        # Build numbering cache upfront — shared across all paragraphs.
        # Avoids re-walking the numbering XML for every list paragraph.
        self._numbering_cache: dict[tuple[int, int], str | None] = {}
        self._numbering_root: Any = None
        try:
            numbering_part = document.part.numbering_part
            if numbering_part is not None:
                self._numbering_root = numbering_part.element
        except (AttributeError, KeyError, NotImplementedError):
            pass

        paragraphs: list[dict[str, Any]] = []
        tables: list[dict[str, Any]] = []
        document_order: list[dict[str, Any]] = []
        paragraph_index = 0
        table_index = 0

        # Single pass over body children — builds document_order, paragraphs,
        # and tables together rather than iterating the document twice.
        for child in document.element.body.iterchildren():
            tag = child.tag.rsplit("}", 1)[-1]

            if tag == "p":
                para_obj = Paragraph(child, document)
                para_dict = self._extract_paragraph(para_obj, paragraph_index)
                paragraphs.append(para_dict)
                document_order.append(
                    {"type": "paragraph", "index": paragraph_index})
                paragraph_index += 1

            elif tag == "tbl":
                tbl_obj = Table(child, document)
                tbl_dict = self._extract_table(tbl_obj, table_index)
                tables.append(tbl_dict)
                document_order.append(
                    {"type": "table", "index": table_index})
                table_index += 1

        self.logger.info(
            "DOCX document structure",
            extra={"paragraphs": paragraph_index, "tables": table_index},
        )

        return {
            "document_order": document_order,
            "document_defaults": None,  # not needed for chunking
            "styles": [],               # not needed for chunking
            "paragraphs": paragraphs,
            "tables": tables,
            "media": [],                # include_media=False for chunking
        }

    def _get_paragraph_list_info(self, paragraph: Paragraph) -> dict[str, Any] | None:
        """Extract list/numbering info from a paragraph element."""
        p_el = paragraph._element  # pylint: disable=protected-access
        p_pr = p_el.pPr
        num_pr = p_pr.numPr if p_pr is not None else None
        if num_pr is None:
            return None
        num_id = num_pr.numId.val if num_pr.numId is not None else None
        ilvl = num_pr.ilvl.val if num_pr.ilvl is not None else None
        return {
            "num_id": int(num_id) if num_id is not None else None,
            "level": int(ilvl) if ilvl is not None else None,
        }

    def _get_list_format_flags(
        self, style_name: str | None, numbering_format: str | None
    ) -> tuple[bool, bool]:
        """Return (is_bullet, is_numbered) derived from style name and numbering format."""
        is_bullet = bool(style_name and "bullet" in style_name.lower())
        is_numbered = bool(style_name and "number" in style_name.lower())
        if numbering_format:
            fmt = numbering_format.split(":", 1)[0].lower()
            if fmt == "bullet":
                is_bullet = True
            else:
                is_numbered = True
        return is_bullet, is_numbered

    def _extract_paragraph(
        self,
        paragraph: Paragraph,
        index: int,
    ) -> dict[str, Any]:
        """Extract paragraph text and list metadata. No run/font data."""
        list_info = self._get_paragraph_list_info(paragraph)
        style_name = paragraph.style.name if paragraph.style else None
        numbering_format = self._resolve_list_formatting(list_info)
        is_bullet, is_numbered = self._get_list_format_flags(
            style_name, numbering_format)

        return {
            "index": index,
            "text": paragraph.text,
            "style": style_name,
            "is_bullet": is_bullet or list_info is not None,
            "is_numbered": is_numbered,
            "list_info": list_info,
            "numbering_format": numbering_format,
            "list_level": list_info.get("level") if list_info else None,
            "runs": [],  # not needed for chunking
        }

    def _find_abstract_num_id(self, root: Any, num_id: int) -> int | None:
        """Return abstractNumId for a given numId, or None if not found."""
        for num in root.findall("w:num", XML_NS):
            if int(num.get(qn("w:numId")) or -1) == num_id:
                abstract_elem = num.find("w:abstractNumId", XML_NS)
                if abstract_elem is None:
                    return None
                abstract_id = int(abstract_elem.get(qn(_W_VAL)) or -1)
                return abstract_id if abstract_id >= 0 else None
        return None

    def _extract_lvl_format(self, abs_num: Any, ilvl: int) -> str | None:
        """Return format string for the matching level within an abstractNum element."""
        for lvl in abs_num.findall("w:lvl", XML_NS):
            if int(lvl.get(qn("w:ilvl")) or -1) != ilvl:
                continue
            num_fmt = lvl.find("w:numFmt", XML_NS)
            lvl_text = lvl.find("w:lvlText", XML_NS)
            if num_fmt is not None and lvl_text is not None:
                fmt = num_fmt.get(qn(_W_VAL))
                text = lvl_text.get(qn(_W_VAL))
                if fmt and text:
                    return f"{fmt}:{text}"
            return None
        return None

    def _find_level_format(self, root: Any, abstract_id: int, ilvl: int) -> str | None:
        """Return format string for a given abstractNumId and level, or None."""
        for abs_num in root.findall("w:abstractNum", XML_NS):
            if int(abs_num.get(qn("w:abstractNumId")) or -1) != abstract_id:
                continue
            return self._extract_lvl_format(abs_num, ilvl)
        return None

    def _resolve_list_formatting(
        self, list_info: dict[str, Any] | None
    ) -> str | None:
        """Resolve abstract numbering format, with per-document cache."""
        if not list_info or self._numbering_root is None:
            return None
        num_id = list_info.get("num_id")
        ilvl = list_info.get("level") or 0
        if num_id is None:
            return None

        cache_key = (num_id, ilvl)
        if cache_key in self._numbering_cache:
            return self._numbering_cache[cache_key]

        try:
            abstract_id = self._find_abstract_num_id(
                self._numbering_root, num_id)
            if abstract_id is None:
                result = None
            else:
                result = self._find_level_format(
                    self._numbering_root, abstract_id, ilvl)
        except (AttributeError, ValueError, TypeError, KeyError):
            result = None

        self._numbering_cache[cache_key] = result
        return result

    # ── Table ─────────────────────────────────────────────────────────────────

    def _extract_table(self, table: Table, index: int) -> dict[str, Any]:
        """Extract table cell text only — no runs, no cell paragraphs."""
        rows = []
        for row_idx, row in enumerate(table.rows):
            cells = []
            for cell in row.cells:
                cell_text = "\n".join(
                    p.text for p in cell.paragraphs if p.text)
                cells.append({
                    "text": cell_text,
                    "paragraphs": [],  # not needed for chunking
                })
            rows.append({"cells": cells, "row_index": row_idx})

        return {
            "index": index,
            "row_count": len(table.rows),
            "column_count": len(table.columns),
            "style": None,  # not needed for chunking
            "rows": rows,
        }
