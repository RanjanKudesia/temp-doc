"""DOCX generation pipeline for temp-doc service."""

import base64
import logging
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_COLOR_INDEX
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Emu
from docx.shared import Inches
from docx.shared import Pt
from docx.shared import RGBColor

from ...schemas.temp_doc_schema import (
    ExtractedData,
    ExtractedDocumentDefaults,
    ExtractedMediaItem,
    ExtractedParagraph,
    ExtractedStyle,
    ExtractedTable,
)

W_VAL_LITERAL = "w:val"


class DocxGenerationPipeline:
    """Generate DOCX from extracted data."""

    def __init__(self) -> None:
        self.logger = logging.getLogger(__name__)

    def run(self, extracted_data: ExtractedData, title: str | None = None) -> bytes:
        """Generate DOCX from extracted data."""
        document = Document()

        if title:
            document.add_heading(title, level=1)

        self._add_extracted_payload(document, extracted_data)

        output = BytesIO()
        document.save(output)
        return output.getvalue()

    def _add_extracted_payload(self, document: Document, extracted_data: ExtractedData) -> None:
        self._apply_document_defaults(
            document, extracted_data.document_defaults)
        self._apply_extracted_styles(document, extracted_data.styles)

        paragraph_by_index = {
            item.index: item for item in extracted_data.paragraphs}
        table_by_index = {item.index: item for item in extracted_data.tables}

        if extracted_data.document_order:
            self._add_items_in_document_order(
                document,
                extracted_data,
                paragraph_by_index,
                table_by_index,
            )
            return

        self._add_items_sorted(document, extracted_data)

    def _add_items_in_document_order(
        self,
        document: Document,
        extracted_data: ExtractedData,
        paragraph_by_index: dict[int, ExtractedParagraph],
        table_by_index: dict[int, ExtractedTable],
    ) -> None:
        """Render paragraphs/tables based on preserved source order."""
        for order_item in extracted_data.document_order:
            if order_item.type == "paragraph":
                paragraph = paragraph_by_index.get(order_item.index)
                if paragraph is not None:
                    self._add_extracted_paragraph(document, paragraph)
            elif order_item.type == "table":
                table = table_by_index.get(order_item.index)
                if table is not None:
                    self._add_extracted_table(document, table)

    def _add_items_sorted(self, document: Document, extracted_data: ExtractedData) -> None:
        """Render paragraphs/tables in index order when source order is absent."""
        for paragraph in sorted(extracted_data.paragraphs, key=lambda item: item.index):
            self._add_extracted_paragraph(document, paragraph)
        for table in sorted(extracted_data.tables, key=lambda item: item.index):
            self._add_extracted_table(document, table)

    def _apply_document_defaults(
        self,
        document: Document,
        defaults: ExtractedDocumentDefaults | None,
    ) -> None:
        """Apply source doc defaults to base styles."""
        if defaults is None:
            return

        for style_name in ("Normal", "Default Paragraph Font"):
            try:
                style_obj = document.styles[style_name]
            except KeyError:
                continue

            if defaults.font_name:
                style_obj.font.name = defaults.font_name
            if defaults.font_size_pt is not None and defaults.font_size_pt > 0:
                style_obj.font.size = Pt(defaults.font_size_pt)
            if defaults.color_rgb:
                try:
                    style_obj.font.color.rgb = self._hex_to_rgb_color(
                        defaults.color_rgb)
                except (ValueError, TypeError):
                    pass

    def _apply_extracted_styles(self, document: Document, styles: list[ExtractedStyle]) -> None:
        """Apply extracted style font defaults so inherited formatting is preserved."""
        for style_data, style_obj in self._iter_style_targets(document, styles):
            font_data = style_data.font
            if font_data is None:
                continue
            self._apply_style_font_overrides(style_obj, font_data)

    def _iter_style_targets(self, document: Document, styles: list[ExtractedStyle]):
        """Yield (style_data, style_obj) for styles that can be applied."""
        for style_data in styles:
            if style_data.font is None:
                continue
            style_obj = self._get_or_create_style(document, style_data)
            if style_obj is None:
                continue
            yield style_data, style_obj

    def _apply_style_font_overrides(self, style_obj, font_data) -> None:
        """Apply extracted font settings to a style object."""
        self._apply_style_name(style_obj, font_data)
        self._apply_style_size(style_obj, font_data)
        self._apply_style_bool(style_obj, "bold", font_data.bold, "b")
        self._apply_style_bool(style_obj, "italic", font_data.italic, "i")
        self._apply_style_bool(style_obj, "underline",
                               font_data.underline, "u")
        self._apply_style_color(style_obj, font_data)
        self._apply_style_highlight(style_obj, font_data)

    def _apply_style_name(self, style_obj, font_data) -> None:
        if font_data.name:
            style_obj.font.name = font_data.name
            return
        self._clear_style_rpr_override(style_obj, "rFonts")

    def _apply_style_size(self, style_obj, font_data) -> None:
        if font_data.size_pt is not None and font_data.size_pt > 0:
            style_obj.font.size = Pt(font_data.size_pt)
            return
        self._clear_style_rpr_override(style_obj, "sz")
        self._clear_style_rpr_override(style_obj, "szCs")

    def _apply_style_bool(self, style_obj, attr: str, value: bool | None, clear_tag: str) -> None:
        if value is None:
            self._clear_style_rpr_override(style_obj, clear_tag)
            return
        setattr(style_obj.font, attr, value)

    def _apply_style_color(self, style_obj, font_data) -> None:
        if not font_data.color_rgb:
            self._clear_style_rpr_override(style_obj, "color")
            return
        try:
            style_obj.font.color.rgb = self._hex_to_rgb_color(
                font_data.color_rgb)
        except (ValueError, TypeError):
            pass

    def _apply_style_highlight(self, style_obj, font_data) -> None:
        if not font_data.highlight_color:
            self._clear_style_rpr_override(style_obj, "highlight")
            return
        try:
            style_obj.font.highlight_color = WD_COLOR_INDEX[font_data.highlight_color.upper(
            )]
        except (KeyError, AttributeError):
            pass

    def _clear_style_rpr_override(self, style_obj, tag_name: str) -> None:
        """Remove direct run-property override from style XML so value can inherit."""
        try:
            style_el = style_obj.element
            rpr = style_el.find(qn("w:rPr"))
            if rpr is None:
                return

            child = rpr.find(qn(f"w:{tag_name}"))
            if child is not None:
                rpr.remove(child)
        except (AttributeError, KeyError, TypeError, ValueError):
            return

    def _get_or_create_style(self, document: Document, style_data: ExtractedStyle):
        style_name = style_data.name
        style_id = style_data.style_id

        if style_name:
            try:
                return document.styles[style_name]
            except KeyError:
                pass

        if style_id:
            try:
                return document.styles[style_id]
            except KeyError:
                pass

        if not style_name:
            return None

        style_type = (style_data.type or "").upper()
        if "PARAGRAPH" in style_type:
            create_type = WD_STYLE_TYPE.PARAGRAPH
        elif "CHARACTER" in style_type:
            create_type = WD_STYLE_TYPE.CHARACTER
        elif "TABLE" in style_type:
            create_type = WD_STYLE_TYPE.TABLE
        else:
            create_type = WD_STYLE_TYPE.PARAGRAPH

        try:
            return document.styles.add_style(style_name, create_type)
        except (AttributeError, KeyError, TypeError, ValueError):
            return None

    def _add_extracted_paragraph(self, document: Document, paragraph_data: ExtractedParagraph) -> None:
        paragraph = self._create_output_paragraph(document, paragraph_data)

        self._populate_output_paragraph(paragraph, paragraph_data)

    def _populate_output_paragraph(self, paragraph, paragraph_data: ExtractedParagraph) -> None:
        """Apply paragraph style, alignment, and runs/text to an existing paragraph."""
        style_name = self._resolve_paragraph_style_name(paragraph_data)
        if style_name:
            try:
                paragraph.style = style_name
            except KeyError:
                pass

        alignment = self._map_alignment(paragraph_data.alignment)
        if alignment is not None:
            paragraph.alignment = alignment

        if paragraph_data.runs:
            self._add_paragraph_runs(paragraph, paragraph_data)
            return
        if paragraph_data.text:
            paragraph.add_run(paragraph_data.text)

    def _create_output_paragraph(self, document: Document, paragraph_data: ExtractedParagraph):
        """Create paragraph with best-effort style assignment."""
        style_name = self._resolve_paragraph_style_name(paragraph_data)
        if style_name:
            try:
                return document.add_paragraph(style=style_name)
            except KeyError:
                return document.add_paragraph()
        return document.add_paragraph()

    def _add_paragraph_runs(self, paragraph, paragraph_data: ExtractedParagraph) -> None:
        """Add runs and embedded media to paragraph."""
        for run_data in paragraph_data.runs:
            if run_data.hyperlink_url:
                self._add_hyperlink_run(paragraph, run_data)
                continue
            run = paragraph.add_run(run_data.text or "")
            self._apply_run_formatting(run, run_data)
            for media_item in run_data.embedded_media:
                self._add_media_to_paragraph(paragraph, media_item)

    def _apply_run_formatting(self, run, run_data) -> None:
        if run_data.bold is not None:
            run.bold = run_data.bold
        if run_data.italic is not None:
            run.italic = run_data.italic
        if run_data.underline is not None:
            run.underline = run_data.underline
        if run_data.font_name:
            run.font.name = run_data.font_name
        if run_data.font_size_pt is not None and run_data.font_size_pt > 0:
            run.font.size = Pt(run_data.font_size_pt)
        if run_data.color_rgb:
            try:
                run.font.color.rgb = self._hex_to_rgb_color(run_data.color_rgb)
            except (ValueError, TypeError):
                pass
        if run_data.highlight_color:
            try:
                run.font.highlight_color = WD_COLOR_INDEX[run_data.highlight_color.upper(
                )]
            except (KeyError, AttributeError):
                pass

    def _add_hyperlink_run(self, paragraph, run_data) -> None:
        url = run_data.hyperlink_url or ""
        text = run_data.text or ""
        try:
            r_id = paragraph.part.relate_to(
                url, RT.HYPERLINK, is_external=True)
        except (AttributeError, KeyError, TypeError, ValueError):
            run = paragraph.add_run(text)
            self._apply_run_formatting(run, run_data)
            return

        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)

        run_elem = self._build_hyperlink_run_element(run_data, text)
        self._append_text_to_oxml_run(run_elem, text)

        hyperlink.append(run_elem)
        # python-docx has no public API to append hyperlink nodes directly.
        # pylint: disable=protected-access
        paragraph._p.append(hyperlink)

    def _build_hyperlink_run_element(self, run_data, text: str):
        """Build oxml run element with hyperlink style overrides."""
        del text
        run_elem = OxmlElement("w:r")
        rpr = OxmlElement("w:rPr")

        hyperlink_blue = self._resolve_hyperlink_color(run_data)
        color_elem = OxmlElement("w:color")
        color_elem.set(qn(W_VAL_LITERAL), hyperlink_blue)
        rpr.append(color_elem)

        if run_data.underline is not False:
            u_elem = OxmlElement("w:u")
            u_elem.set(qn(W_VAL_LITERAL), "single")
            rpr.append(u_elem)

        if run_data.bold:
            rpr.append(OxmlElement("w:b"))
        if run_data.italic:
            rpr.append(OxmlElement("w:i"))
        if run_data.font_name:
            r_fonts = OxmlElement("w:rFonts")
            r_fonts.set(qn("w:ascii"), run_data.font_name)
            r_fonts.set(qn("w:hAnsi"), run_data.font_name)
            rpr.append(r_fonts)
        if run_data.font_size_pt and run_data.font_size_pt > 0:
            half_pts = str(int(run_data.font_size_pt * 2))
            sz = OxmlElement("w:sz")
            sz.set(qn(W_VAL_LITERAL), half_pts)
            rpr.append(sz)
            sz_cs = OxmlElement("w:szCs")
            sz_cs.set(qn(W_VAL_LITERAL), half_pts)
            rpr.append(sz_cs)

        run_elem.append(rpr)
        return run_elem

    def _resolve_hyperlink_color(self, run_data) -> str:
        """Return hyperlink color hex (without #), defaulting to theme-friendly blue."""
        if not run_data.color_rgb:
            return "0563C1"
        return run_data.color_rgb.replace("#", "").strip()

    def _append_text_to_oxml_run(self, run_elem, text: str) -> None:
        buf: list[str] = []

        def flush_text() -> None:
            if not buf:
                return
            content = "".join(buf)
            buf.clear()
            t = OxmlElement("w:t")
            t.text = content
            if content.startswith(" ") or content.endswith(" "):
                t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            run_elem.append(t)

        for ch in text:
            if ch == "\n":
                flush_text()
                run_elem.append(OxmlElement("w:br"))
            elif ch == "\t":
                flush_text()
                run_elem.append(OxmlElement("w:tab"))
            else:
                buf.append(ch)

        flush_text()

    def _add_extracted_table(self, document: Document, table_data: ExtractedTable) -> None:
        if not table_data.rows:
            return

        column_count = max((len(row.cells)
                           for row in table_data.rows), default=0)
        if column_count == 0:
            return

        table = document.add_table(
            rows=len(table_data.rows), cols=column_count)
        self._apply_table_style(table, table_data)
        self._populate_docx_table(table, table_data)

    def _apply_table_style(self, table, table_data: ExtractedTable) -> None:
        """Apply extracted table style when available."""
        if not table_data.style:
            return

        try:
            table.style = table_data.style
        except KeyError:
            return

    def _populate_docx_table(self, table, table_data: ExtractedTable) -> None:
        """Populate a docx table recursively from extracted table data."""
        for row_index, row in enumerate(table_data.rows):
            for column_index, cell_data in enumerate(row.cells):
                target_cell = table.cell(row_index, column_index)
                merged_cell = self._merge_target_cell(
                    table,
                    target_cell,
                    cell_data,
                    row_index,
                    column_index,
                )
                self._populate_docx_cell(merged_cell, cell_data)

    def _merge_target_cell(self, table, target_cell, cell_data, row_index: int, column_index: int):
        """Merge table cells when colspan or rowspan is present."""
        colspan = getattr(cell_data, "colspan", None) or 1
        rowspan = getattr(cell_data, "rowspan", None) or 1
        if colspan == 1 and rowspan == 1:
            return target_cell

        end_row = min(row_index + rowspan - 1, len(table.rows) - 1)
        end_col = min(column_index + colspan - 1, len(table.columns) - 1)
        if end_row == row_index and end_col == column_index:
            return target_cell

        return target_cell.merge(table.cell(end_row, end_col))

    def _populate_docx_cell(self, cell, cell_data) -> None:
        """Populate a docx table cell with paragraphs and nested tables."""
        cell.text = ""
        paragraphs = list(getattr(cell_data, "paragraphs", []) or [])

        if paragraphs:
            self._populate_existing_cell_paragraph(
                cell.paragraphs[0], paragraphs[0])
            for paragraph_data in paragraphs[1:]:
                paragraph = cell.add_paragraph()
                self._populate_output_paragraph(paragraph, paragraph_data)
        elif getattr(cell_data, "text", None):
            cell.paragraphs[0].add_run(cell_data.text)

        for nested_table in getattr(cell_data, "tables", []) or []:
            self._add_nested_docx_table(cell, nested_table)

    def _populate_existing_cell_paragraph(self, paragraph, paragraph_data: ExtractedParagraph) -> None:
        """Populate the default paragraph already present in a table cell."""
        paragraph.text = ""
        self._populate_output_paragraph(paragraph, paragraph_data)

    def _add_nested_docx_table(self, cell, table_data: ExtractedTable) -> None:
        """Add a nested table inside a table cell."""
        if not table_data.rows:
            return

        column_count = max((len(row.cells)
                           for row in table_data.rows), default=0)
        if column_count == 0:
            return

        nested_table = cell.add_table(
            rows=len(table_data.rows), cols=column_count)
        self._apply_table_style(nested_table, table_data)
        self._populate_docx_table(nested_table, table_data)

    def _resolve_paragraph_style_name(self, paragraph_data: ExtractedParagraph) -> str | None:
        if paragraph_data.is_numbered:
            return "List Number"

        if paragraph_data.is_bullet:
            return "List Bullet"

        if paragraph_data.numbering_format:
            fmt = paragraph_data.numbering_format.split(":", 1)[0].lower()
            if fmt == "bullet":
                return "List Bullet"
            return "List Number"

        if paragraph_data.style:
            return paragraph_data.style

        return None

    def _map_alignment(self, raw_alignment: str | None) -> WD_ALIGN_PARAGRAPH | None:
        if raw_alignment is None:
            return None

        normalized = raw_alignment.strip().upper()
        if normalized.startswith("LEFT"):
            return WD_ALIGN_PARAGRAPH.LEFT
        if normalized.startswith("CENTER"):
            return WD_ALIGN_PARAGRAPH.CENTER
        if normalized.startswith("RIGHT"):
            return WD_ALIGN_PARAGRAPH.RIGHT
        if normalized.startswith("JUSTIFY"):
            return WD_ALIGN_PARAGRAPH.JUSTIFY
        return None

    def _add_media_to_paragraph(self, paragraph, media_item: ExtractedMediaItem | str) -> None:
        """Insert an inline image using file path or in-payload base64 bytes."""
        local_file_path = (
            media_item.local_file_path if hasattr(
                media_item, "local_file_path") else media_item
        )
        base64_data = getattr(media_item, "base64_data", None) or getattr(
            media_item, "base64", None)

        width_emu = getattr(media_item, "width_emu", None)
        height_emu = getattr(media_item, "height_emu", None)

        try:
            run = paragraph.add_run()

            picture_source = None
            if base64_data:
                picture_source = BytesIO(base64.b64decode(base64_data))
            elif local_file_path:
                media_path = Path(local_file_path)
                if media_path.exists() and media_path.is_file():
                    picture_source = str(media_path)

            if picture_source is None:
                return

            if width_emu and height_emu:
                run.add_picture(picture_source, width=Emu(
                    width_emu), height=Emu(height_emu))
            elif width_emu:
                run.add_picture(picture_source, width=Emu(width_emu))
            else:
                run.add_picture(picture_source, width=Inches(2.5))
        except (
            AttributeError,
            OSError,
            TypeError,
            ValueError,
        ):
            return

    def _hex_to_rgb_color(self, value: str) -> RGBColor:
        hex_str = value.replace("#", "").strip()
        return RGBColor(
            int(hex_str[0:2], 16),
            int(hex_str[2:4], 16),
            int(hex_str[4:6], 16),
        )
