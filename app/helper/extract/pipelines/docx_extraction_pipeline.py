"""DOCX extraction pipeline for temp-doc service - simplified, no storage."""

import base64
import json
import logging
import mimetypes
import time
from io import BytesIO
from typing import Any
from zipfile import is_zipfile

from docx import Document
from docx.document import Document as DocumentObject
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from lxml import etree

XML_NS = {
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
}

_W_VAL = "w:val"


class DocxExtractionPipeline:
    """Extract paragraphs, tables, and styles from DOCX files."""

    def __init__(self) -> None:
        self.logger = logging.getLogger(__name__)

    def run(self, file_bytes: bytes, include_media: bool = True) -> dict[str, Any]:
        """Parse a DOCX byte stream and return extracted JSON payload."""
        t0 = time.perf_counter()
        self.logger.info(
            "DOCX extraction started",
            extra={"file_size_bytes": len(file_bytes)},
        )

        if not is_zipfile(BytesIO(file_bytes)):
            raise ValueError(
                "Invalid DOCX file: not a valid ZIP archive. File may be corrupted.")

        try:
            document = Document(BytesIO(file_bytes))
        except (ValueError, TypeError, OSError, KeyError) as e:
            raise ValueError(
                f"Failed to parse DOCX document: {str(e)}") from e

        extracted = self._extract_document(
            document, include_media=include_media)

        elapsed_ms = round((time.perf_counter() - t0) * 1000)
        response_size_bytes = len(json.dumps(extracted).encode("utf-8"))
        self.logger.info(
            "DOCX extraction complete",
            extra={
                "elapsed_ms": elapsed_ms,
                "elapsed_s": round(elapsed_ms / 1000, 3),
                "response_size_bytes": response_size_bytes,
                "response_size_kb": round(response_size_bytes / 1024, 2),
                "paragraphs_extracted": len(extracted.get("paragraphs", [])),
                "tables_extracted": len(extracted.get("tables", [])),
                "media_extracted": len(extracted.get("media", [])),
                "styles_extracted": len(extracted.get("styles", [])),
            },
        )
        return extracted

    def _extract_document(
        self,
        document: DocumentObject,
        include_media: bool,
    ) -> dict[str, Any]:
        """Extract all document content."""
        total_paragraphs = len(document.paragraphs)
        total_tables = len(document.tables)
        total_sections = len(document.sections)
        total_styles = len(document.styles)
        total_shapes = len(document.inline_shapes)

        self.logger.info(
            "DOCX document structure",
            extra={
                "paragraphs": total_paragraphs,
                "tables": total_tables,
                "sections": total_sections,
                "styles": total_styles,
                "inline_shapes": total_shapes,
            },
        )

        media_index = (
            self._extract_and_save_media(document, "temp-doc")
            if include_media
            else {}
        )

        paragraphs: list[dict[str, Any]] = []
        log_interval = max(1, total_paragraphs //
                           10) if total_paragraphs >= 10 else total_paragraphs
        for index, paragraph in enumerate(document.paragraphs):
            paragraphs.append(
                self._extract_paragraph(
                    paragraph, index, document, media_index)
            )
            if total_paragraphs > 0 and ((index + 1) % log_interval == 0 or index + 1 == total_paragraphs):
                self.logger.debug(
                    "Paragraphs extracted",
                    extra={"done": index + 1, "total": total_paragraphs},
                )

        tables: list[dict[str, Any]] = []
        for index, table in enumerate(document.tables):
            tables.append(
                self._extract_table(table, index, document, media_index)
            )
            self.logger.debug(
                "Table extracted",
                extra={
                    "table_index": index,
                    "rows": len(table.rows),
                    "columns": len(table.columns) if table.rows else 0,
                },
            )

        body_order: list[dict[str, Any]] = []
        paragraph_index = 0
        table_index = 0
        for child in document.element.body.iterchildren():
            tag = child.tag.rsplit("}", 1)[-1]
            if tag == "p":
                body_order.append(
                    {"type": "paragraph", "index": paragraph_index})
                paragraph_index += 1
            elif tag == "tbl":
                body_order.append({"type": "table", "index": table_index})
                table_index += 1

        styles = self._extract_styles(document)
        document_defaults = self._extract_document_defaults(document)

        return {
            "document_order": body_order,
            "document_defaults": document_defaults,
            "styles": styles,
            "paragraphs": paragraphs,
            "tables": tables,
            "media": list(media_index.values()),
        }

    def _extract_and_save_media(
        self, document: DocumentObject, output_basename: str
    ) -> dict[str, dict[str, Any]]:
        """Extract embedded images and keep them in JSON as base64 payload."""
        media_index: dict[str, dict[str, Any]] = {}

        for rel_id, rel in document.part.rels.items():
            if rel.reltype != RT.IMAGE:
                continue

            try:
                image_part = rel.target_part
                if image_part is None or not hasattr(image_part, "blob"):
                    continue

                blob = image_part.blob
                if not blob:
                    continue

                content_type = image_part.content_type or "application/octet-stream"
                extension = self._content_type_to_extension(content_type)
                file_name = f"{output_basename}_{rel_id}.{extension}"
                media_index[rel_id] = {
                    "relationship_id": rel_id,
                    "content_type": content_type,
                    "file_name": file_name,
                    "local_file_path": None,
                    "local_url": None,
                    "base64": base64.b64encode(blob).decode("ascii"),
                    "size_bytes": len(blob),
                    "source_partname": str(image_part.partname),
                    "alt_text": None,
                }
            except (AttributeError, KeyError, ValueError, TypeError, OSError):
                continue

        self.logger.info(
            "Media extraction complete",
            extra={"images_found": len(media_index)},
        )
        return media_index

    def _content_type_to_extension(self, content_type: str) -> str:
        normalized = content_type.lower().strip()
        if not normalized:
            return "bin"

        ext = mimetypes.guess_extension(normalized, strict=False)
        if ext:
            return ext.lstrip(".")

        fallback = {
            "image/jpeg": "jpg",
            "image/jpg": "jpg",
            "image/png": "png",
            "image/gif": "gif",
            "image/bmp": "bmp",
            "image/tiff": "tiff",
            "image/webp": "webp",
            "image/svg+xml": "svg",
            "image/x-wmf": "wmf",
            "image/x-emf": "emf",
        }
        return fallback.get(normalized, "bin")

    def _extract_document_defaults(self, document: DocumentObject) -> dict[str, Any]:
        """Extract DOCX-level default run properties."""
        defaults = {
            "font_name": None,
            "font_size_pt": None,
            "color_rgb": None,
        }

        try:
            styles_root = document.styles.element
            run_defaults = styles_root.find(
                "w:docDefaults/w:rPrDefault/w:rPr", XML_NS)
            if run_defaults is None:
                return defaults

            r_fonts = run_defaults.find("w:rFonts", XML_NS)
            if r_fonts is not None:
                font_name = (
                    r_fonts.get(qn("w:ascii"))
                    or r_fonts.get(qn("w:hAnsi"))
                    or r_fonts.get(qn("w:cs"))
                )
                if font_name:
                    defaults["font_name"] = font_name

            sz_elem = run_defaults.find("w:sz", XML_NS)
            if sz_elem is not None:
                sz_val = sz_elem.get(qn(_W_VAL))
                if sz_val is not None:
                    defaults["font_size_pt"] = int(sz_val) / 2.0

            color_elem = run_defaults.find("w:color", XML_NS)
            if color_elem is not None:
                color_val = color_elem.get(qn(_W_VAL))
                if color_val and color_val.lower() != "auto":
                    defaults["color_rgb"] = color_val.upper()

        # type: ignore[attr-defined]
        except (
            AttributeError, KeyError, ValueError, TypeError,
            # pylint: disable=c-extension-no-member
            etree.XMLSyntaxError,
        ) as e:
            self.logger.debug(
                "Failed to extract document defaults: %s", str(e))

        return defaults

    def _extract_styles(self, document: DocumentObject) -> list[dict[str, Any]]:
        """Extract paragraph and character styles."""
        styles = []
        try:
            for style in document.styles:
                style_dict = {
                    "style_id": style.style_id,
                    "name": style.name,
                    "type": str(style.type) if style.type else None,
                    "font": None,
                }

                # Try to extract font info if it's a paragraph or character style
                try:
                    font = style.font
                    if font:
                        style_dict["font"] = {
                            "name": font.name,
                            "size_pt": font.size.pt if font.size else None,
                            "bold": font.bold,
                            "italic": font.italic,
                            "underline": font.underline,
                            "color_rgb": str(
                                font.color.rgb) if font.color and font.color.rgb else None,
                        }
                except (AttributeError, TypeError, ValueError, KeyError):
                    pass

                styles.append(style_dict)
        except (AttributeError, KeyError, ValueError, TypeError, OSError) as e:
            self.logger.debug("Failed to extract styles: %s", str(e))

        return styles

    def _get_paragraph_list_info(self, paragraph: Paragraph) -> dict[str, Any] | None:
        """Extract list/numbering info from a paragraph element."""
        p_el = paragraph._element  # pylint: disable=protected-access
        p_pr = p_el.pPr
        num_pr = p_pr.numPr if p_pr is not None else None
        if num_pr is None:
            return None
        num_id = num_pr.numId.val if num_pr.numId is not None else None
        ilvl = num_pr.ilvl.val if num_pr.ilvl is not None else None
        return {
            "num_id": int(num_id) if num_id is not None else None,
            "level": int(ilvl) if ilvl is not None else None,
        }

    def _get_list_format_flags(
        self, style_name: str | None, numbering_format: str | None
    ) -> tuple[bool, bool]:
        """Return (is_bullet, is_numbered) derived from style name and numbering format."""
        is_bullet = bool(style_name and "bullet" in style_name.lower())
        is_numbered = bool(style_name and "number" in style_name.lower())
        if numbering_format:
            fmt = numbering_format.split(":", 1)[0].lower()
            if fmt == "bullet":
                is_bullet = True
            else:
                is_numbered = True
        return is_bullet, is_numbered

    def _extract_paragraph(
        self,
        paragraph: Paragraph,
        index: int,
        document: DocumentObject,
        media_index: dict[str, dict[str, Any]],
    ) -> dict[str, Any]:
        """Extract paragraph with formatting."""
        runs = [self._extract_run(run, i, media_index)
                for i, run in enumerate(paragraph.runs)]

        alignment = paragraph.alignment.name if paragraph.alignment is not None else None
        list_info = self._get_paragraph_list_info(paragraph)
        style_name = paragraph.style.name if paragraph.style else None
        numbering_format = self._resolve_list_formatting(list_info, document)
        is_bullet, is_numbered = self._get_list_format_flags(
            style_name, numbering_format)

        return {
            "index": index,
            "text": paragraph.text,
            "style": style_name,
            "alignment": alignment,
            "is_bullet": is_bullet or list_info is not None,
            "is_numbered": is_numbered,
            "list_info": list_info,
            "numbering_format": numbering_format,
            "list_level": list_info.get("level") if list_info else None,
            "runs": runs,
        }

    def _find_abstract_num_id(self, root: Any, num_id: int) -> int | None:
        """Return abstractNumId for a given numId, or None if not found."""
        for num in root.findall("w:num", XML_NS):
            if int(num.get(qn("w:numId")) or -1) == num_id:
                abstract_elem = num.find("w:abstractNumId", XML_NS)
                if abstract_elem is None:
                    return None
                abstract_id = int(abstract_elem.get(qn(_W_VAL)) or -1)
                return abstract_id if abstract_id >= 0 else None
        return None

    def _extract_lvl_format(self, abs_num: Any, ilvl: int) -> str | None:
        """Return format string for the matching level within an abstractNum element."""
        for lvl in abs_num.findall("w:lvl", XML_NS):
            if int(lvl.get(qn("w:ilvl")) or -1) != ilvl:
                continue
            num_fmt = lvl.find("w:numFmt", XML_NS)
            lvl_text = lvl.find("w:lvlText", XML_NS)
            if num_fmt is not None and lvl_text is not None:
                fmt = num_fmt.get(qn(_W_VAL))
                text = lvl_text.get(qn(_W_VAL))
                if fmt and text:
                    return f"{fmt}:{text}"
            return None
        return None

    def _find_level_format(self, root: Any, abstract_id: int, ilvl: int) -> str | None:
        """Return format string for a given abstractNumId and level, or None."""
        for abs_num in root.findall("w:abstractNum", XML_NS):
            if int(abs_num.get(qn("w:abstractNumId")) or -1) != abstract_id:
                continue
            return self._extract_lvl_format(abs_num, ilvl)
        return None

    def _resolve_list_formatting(
        self, list_info: dict[str, Any] | None, document: DocumentObject
    ) -> str | None:
        """Resolve abstract numbering format to human-readable list format."""
        if not list_info:
            return None
        num_id = list_info.get("num_id")
        ilvl = list_info.get("level") or 0
        if num_id is None:
            return None
        try:
            numbering_part = document.part.numbering_part
            if numbering_part is None:
                return None
            root = numbering_part.element
            abstract_id = self._find_abstract_num_id(root, num_id)
            if abstract_id is None:
                return None
            return self._find_level_format(root, abstract_id, ilvl)
        except (AttributeError, ValueError, TypeError, KeyError):
            return None

    def _extract_run(
        self, run: Run, index: int, media_index: dict[str, dict[str, Any]]
    ) -> dict[str, Any]:
        """Extract text run with formatting."""
        text_color = None
        if run.font.color and run.font.color.rgb:
            text_color = str(run.font.color.rgb)

        embedded_media = []
        run_el = run._element  # pylint: disable=protected-access
        drawing_blips = run_el.xpath(".//*[local-name()='blip']")
        for blip in drawing_blips:
            rel_id = blip.get(qn("r:embed"))
            if rel_id and rel_id in media_index:
                embedded_media.append(dict(media_index[rel_id]))

        return {
            "index": index,
            "text": run.text,
            "bold": run.font.bold,
            "italic": run.font.italic,
            "underline": run.font.underline,
            "font_name": run.font.name,
            "font_size_pt": run.font.size.pt if run.font.size else None,
            "color_rgb": text_color,
            "embedded_media": embedded_media,
        }

    def _extract_table(
        self,
        table: Table,
        index: int,
        document: DocumentObject,
        media_index: dict[str, dict[str, Any]],
    ) -> dict[str, Any]:
        """Extract table with cells and content."""
        rows = []
        for row_idx, row in enumerate(table.rows):
            cells = []
            for cell in row.cells:
                cell_text = "\n".join(
                    [p.text for p in cell.paragraphs if p.text])

                cells.append({
                    "text": cell_text,
                    "paragraphs": [
                        self._extract_paragraph(p, i, document, media_index)
                        for i, p in enumerate(cell.paragraphs)
                    ],
                })

            rows.append({"cells": cells, "row_index": row_idx})

        return {
            "index": index,
            "row_count": len(table.rows),
            "column_count": len(table.columns),
            "style": table.style.name if table.style else None,
            "rows": rows,
        }
